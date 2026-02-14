import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
import io
import pandas as pd
import sqlite3
from datetime import datetime
import time
import os
from PIL import Image
import tempfile
import json
import re

# ==========================================
# 1. KONFIGURASI HALAMAN
# ==========================================
st.set_page_config(page_title="Aplikasi RHK PKH Pro 2.0", layout="wide")

# ==========================================
# 2. DEFINISI CONFIG (GLOBAL)
# ==========================================
CONFIG_LAPORAN = {
    "RHK 1 – Laporan Penyaluran bansos": [
        "Melakukan edukasi dan sosialisasi pencairan secara tunai dan non tunai",
        "Melaksanakan Supervisi Permasalahan Bantuan Sosial",
        "Melaksanakan Monitoring/Pemantauan Penyaluran Bantuan Sosial",
        "Melaksanakan Penelitian penyaluran bantuan Sosial"
    ],
    "RHK 2 – Laporan pertemuan P2K2": [
        "Melaksanakan Pertemuan Peningkatan Kemampuan Keluarga (P2K2)"
    ],
    "RHK 3 – Laporan Verifikasi Komitmen dan Pendampingan KPM": [
        "Melaksanakan Verifikasi Komitmen Pendidikan,Kesehatan dan Kesejahteraan Sosial",
        "Melakukan pendampingan, mediasi, dan fasilitasi kepada KPM PKH dalam proses perubahan perilaku, pola pikir yang mandiri dan produktif"
    ],
    "RHK 4 – Rekapitulasi Data KPM graduasi": [
        "Melakukan usulan KPM Graduasi mandiri dan Pemberdayaan PPSE"
    ],
    "RHK 5 – Laporan Data Verifikasi, Validasi dan pemutakhiran data KPM": [
        "Melaksanakan Pemutakhiran Data",
        "Melaksanakan proses bisnis PKH yang meliputi verifikasi validasi calon penerima bantuan sosial"
    ],
    "RHK 6 – persentase penyelesaian laporan kasus adaptif": [
        "Melaksanakan Respon Kasus/Pengaduan/kebencanaan/Kerentanan"
    ],
    "RHK 7 – Laporan Bulanan ASN PPPK": [
        "Membuat laporan bulanan pelaksanaan PKH dan laporan lainnya."
    ],
    "RHK 8 – laporan pelaksana Tugas direktif": [
        "Melaksanakan Tindak Lanjut Hasil Pemeriksaan (TLHP)",
        "Melakukan sosialisasi kebijakan dan bisnis proses PKH kepada aparat pemerintah tingkat kecamatan, desa/ kelurahan, KPM PKH, dan masyarakat umum secara berkala melalui Pertemuan atau media sosial dll",
        "Mengikuti Rapat Koordinasi, Sosialisasi Kebijakan Proses Bisnis PKH dan Penguatan Kapasitas SDM.",
        "Tugas Lainnya (Penugasan lainnya program Kementrian Sosial)"
    ],
    "RHK 9 – Presentase Penyelesaian Penugasan Direktif Pimpinan": [
        "Berperan aktif dalam memanfaatkan, menggunakan, melibatkan dan menyebarkan Media Sosial untuk menyampaikan semua program di Kementerian Sosial"
    ]
}

DAFTAR_USER = {
    "admin": "admin123",
    "pendamping": "pkh2026",
    "user": "user"
}

# --- RESET SESSION JIKA CONFIG BERUBAH ---
if 'selected_rhk' in st.session_state:
    if st.session_state['selected_rhk'] is not None:
        if st.session_state['selected_rhk'] not in CONFIG_LAPORAN:
            st.session_state.clear()
            st.rerun()

# ==========================================
# 3. HELPER FUNCTIONS (IMAGE, TEXT, DB)
# ==========================================

def init_db():
    conn = sqlite3.connect('rhk_pro_fixed.db')
    c = conn.cursor()
    
    # 1. Tabel Users
    c.execute('''CREATE TABLE IF NOT EXISTS users (
        username TEXT PRIMARY KEY, 
        password TEXT
    )''')
    
    # 2. Tabel User Settings
    c.execute('''CREATE TABLE IF NOT EXISTS user_settings (
        id INTEGER PRIMARY KEY AUTOINCREMENT, 
        username TEXT,
        nama TEXT, nip TEXT, kpm INTEGER, 
        prov TEXT, kab TEXT, kec TEXT, kel TEXT, jabatan TEXT
    )''')
    
    try:
        c.execute("ALTER TABLE user_settings ADD COLUMN username TEXT")
    except: pass
    try:
        c.execute("ALTER TABLE user_settings ADD COLUMN jabatan TEXT")
    except: pass

    # 3. Tabel Riwayat Laporan
    c.execute('''CREATE TABLE IF NOT EXISTS riwayat_laporan (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT,
        tanggal TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        bulan TEXT,
        tahun TEXT,
        jenis_rhk TEXT,
        judul_kegiatan TEXT,
        file_docx BLOB,
        file_pdf BLOB
    )''')
    
    try:
        c.execute("ALTER TABLE riwayat_laporan ADD COLUMN username TEXT")
    except: pass
    
    conn.commit()
    conn.close()

def get_user_settings(username):
    conn = sqlite3.connect('rhk_pro_fixed.db'); c = conn.cursor()
    c.execute('SELECT nama, nip, kpm, prov, kab, kec, kel, jabatan FROM user_settings WHERE username=?', (username,))
    data = c.fetchone(); conn.close()
    
    if data and len(data) == 8: return data
    return ("Nama User", "NIP", 0, "Provinsi", "Kabupaten", "Kecamatan", "Kelurahan", "Pendamping Sosial")

def save_user_settings(username, nama, nip, kpm, prov, kab, kec, kel, jabatan):
    conn = sqlite3.connect('rhk_pro_fixed.db'); c = conn.cursor()
    c.execute("SELECT id FROM user_settings WHERE username=?", (username,))
    exists = c.fetchone()
    
    if exists:
        c.execute('''UPDATE user_settings SET nama=?, nip=?, kpm=?, prov=?, kab=?, kec=?, kel=?, jabatan=? WHERE username=?''', 
                  (nama, nip, kpm, prov, kab, kec, kel, jabatan, username))
    else:
        c.execute('''INSERT INTO user_settings (username, nama, nip, kpm, prov, kab, kec, kel, jabatan) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)''', 
                  (username, nama, nip, kpm, prov, kab, kec, kel, jabatan))
    conn.commit(); conn.close()

def save_to_history(username, bulan, tahun, rhk, judul, docx_io, pdf_io):
    conn = sqlite3.connect('rhk_pro_fixed.db')
    c = conn.cursor()
    docx_io.seek(0); docx_blob = docx_io.getvalue()
    if isinstance(pdf_io, bytes): pdf_blob = pdf_io
    else: pdf_io.seek(0); pdf_blob = pdf_io.getvalue()
    c.execute('''INSERT INTO riwayat_laporan (username, bulan, tahun, jenis_rhk, judul_kegiatan, file_docx, file_pdf) VALUES (?, ?, ?, ?, ?, ?, ?)''', 
              (username, bulan, tahun, rhk, judul, docx_blob, pdf_blob))
    conn.commit(); conn.close()

def get_history_by_filter(username, bulan, tahun):
    conn = sqlite3.connect('rhk_pro_fixed.db'); conn.row_factory = sqlite3.Row; c = conn.cursor()
    if bulan == "SEMUA" and tahun == "SEMUA": 
        c.execute("SELECT id, tanggal, bulan, tahun, jenis_rhk, judul_kegiatan FROM riwayat_laporan WHERE username=? ORDER BY id DESC", (username,))
    elif bulan == "SEMUA": 
        c.execute("SELECT id, tanggal, bulan, tahun, jenis_rhk, judul_kegiatan FROM riwayat_laporan WHERE username=? AND tahun=? ORDER BY id DESC", (username, tahun))
    else: 
        c.execute("SELECT id, tanggal, bulan, tahun, jenis_rhk, judul_kegiatan FROM riwayat_laporan WHERE username=? AND bulan=? AND tahun=? ORDER BY id DESC", (username, bulan, tahun))
    data = c.fetchall(); conn.close(); return data

def get_file_from_history(id_laporan, tipe='docx'):
    conn = sqlite3.connect('rhk_pro_fixed.db'); c = conn.cursor()
    col = 'file_docx' if tipe == 'docx' else 'file_pdf'
    c.execute(f"SELECT {col}, judul_kegiatan FROM riwayat_laporan WHERE id=?", (id_laporan,))
    data = c.fetchone(); conn.close()
    if data: return data[0], data[1]
    return None, None

def delete_history_item(id_laporan):
    conn = sqlite3.connect('rhk_pro_fixed.db'); c = conn.cursor()
    c.execute("DELETE FROM riwayat_laporan WHERE id=?", (id_laporan,)); conn.commit(); conn.close()

def compress_image(uploaded_file, quality=60, max_width=600):
    try:
        uploaded_file.seek(0)
        image = Image.open(uploaded_file)
        if image.mode in ("RGBA", "P"): image = image.convert("RGB")
        if image.width > max_width:
            ratio = max_width / float(image.width)
            new_height = int((float(image.height) * float(ratio)))
            image = image.resize((max_width, new_height), Image.Resampling.LANCZOS)
        output = io.BytesIO()
        image.save(output, format="JPEG", quality=quality, optimize=True)
        output.seek(0); uploaded_file.seek(0)
        return output
    except: uploaded_file.seek(0); return uploaded_file

def clean_text_for_pdf(text):
    if text is None: return "-" 
    text = str(text) 
    text = text.replace('\u2013', '-').replace('\u201c', '"').replace('\u201d', '"')
    return text.encode('latin-1', 'replace').decode('latin-1')

# ==========================================
# 4. GENERATOR AI & DOKUMEN (UPDATED)
# ==========================================

try:
    GOOGLE_API_KEY = st.secrets["GOOGLE_API_KEY"]
    genai.configure(api_key=GOOGLE_API_KEY)
    model = genai.GenerativeModel('gemini-2.0-flash')
except Exception as e:
    st.error(f"Gagal konfigurasi AI: {e}")

def generate_isi_laporan(topik, detail, kpm_total, kpm_fokus, bulan, lokasi_lengkap, ket_info=""):
    # Pastikan ket_info tidak kosong agar AI tidak halusinasi
    if not ket_info or str(ket_info).strip() == "" or str(ket_info).strip() == "-":
        ket_info_prompt = "Kegiatan dilaksanakan sesuai prosedur operasional standar (SOP) yang berlaku."
    else:
        ket_info_prompt = ket_info

    prompt = f"""
    Role: Pendamping Sosial PKH Profesional & Berpengalaman Kemensos RI.
    Tugas: Buat konten Laporan Kegiatan Bulanan yang formal, baku, dan administratif.
    
    Data Laporan:
    - Topik RHK: {topik}
    - Judul Standar: {detail}
    - Lokasi: {lokasi_lengkap}
    - Bulan: {bulan}
    - Sasaran: {kpm_fokus} ({kpm_total} orang)

    === INSTRUKSI UTAMA (PENTING) ===
    INPUT KEGIATAN DARI USER: "{ket_info_prompt}"
    
    1. Pada bagian JSON "kegiatan": Anda WAJIB mengembangkan narasi kronologis BERDASARKAN "INPUT KEGIATAN DARI USER" di atas. 
    2. Gunakan "INPUT KEGIATAN DARI USER" sebagai inti cerita. Jangan membuat cerita yang melenceng dari input tersebut.
    3. Jika input user singkat (misal: "di rumah ketua kelompok mawar"), kembangkan kalimatnya menjadi formal (misal: "Kegiatan dilaksanakan bertempat di kediaman Ketua Kelompok Mawar...").
    
    Instruksi Bagian "Dasar Hukum":
    JANGAN PERNAH menggunakan kalimat "Surat Tugas dari Koordinator PKH Kabupaten/Kota Nomor: ... tanggal ...".
    Gunakan HANYA poin-poin dasar hukum normatif berikut ini:
    1. Peraturan Menteri Sosial Republik Indonesia Nomor 1 Tahun 2018 tentang Program Keluarga Harapan.
    2. Keputusan Direktur Jaminan Sosial Keluarga tentang Petunjuk Teknis Pelaksanaan Program Keluarga Harapan yang berlaku.
    3. Rencana Kerja Tahunan (RKT) Pendamping Sosial PKH Tahun 2026.

    Instruksi Format:
    Gunakan Bahasa Indonesia Ejaan Yang Disempurnakan (EYD) yang baku, kalimat efektif, dan gaya bahasa laporan resmi dinas sosial.

    Output Wajib JSON (tanpa markdown ```json):
    {{
      "pendahuluan": {{
         "umum": "Paragraf pembuka yang menjelaskan latar belakang umum program PKH dan pentingnya kegiatan ini.",
         "maksud_tujuan": "Paragraf menjelaskan maksud dan tujuan spesifik dilaksanakannya kegiatan {detail}.",
         "ruang_lingkup": "Paragraf menjelaskan batasan kegiatan, lokasi, dan sasaran peserta.",
         "dasar": ["Peraturan Menteri Sosial Republik Indonesia Nomor 1 Tahun 2018 tentang Program Keluarga Harapan", "Keputusan Direktur Jaminan Sosial Keluarga tentang Petunjuk Teknis Pelaksanaan Program Keluarga Harapan", "Rencana Kerja Tahunan (RKT) Pendamping Sosial PKH Tahun 2026"]
      }},
      "kegiatan": "Paragraf narasi detail yang menjelaskan jalannya kegiatan. INGAT: Kembangkan paragraf ini berdasarkan '{ket_info_prompt}'. Jelaskan proses dari awal pembukaan hingga penutup sesuai konteks input user tersebut.",
      "hasil": "Paragraf atau poin-poin yang menjelaskan output konkret, dampak, atau hasil yang dicapai dari kegiatan tersebut.",
      "simpulan_saran": "Paragraf berisi kesimpulan evaluatif dan saran konstruktif untuk perbaikan ke depan.",
      "penutup": "Kalimat penutup laporan yang formal."
    }}
    """
    try:
        response = model.generate_content(prompt)
        text = response.text.replace("```json", "").replace("```", "").strip()
        return json.loads(text)
    except Exception as e: return None

def create_word_doc(data, meta, imgs, kop, ttd, extra_info=None, kpm_data=None):
    if data is None: return None
    doc = Document()
    
    # Margin Standar Dinas
    for s in doc.sections: 
        s.top_margin=Cm(2.54)
        s.bottom_margin=Cm(2.54)
        s.left_margin=Cm(2.54)
        s.right_margin=Cm(2.54)

    # --- KOP SURAT ---
    if kop: 
        try: 
            p = doc.add_paragraph()
            p.alignment = 1
            # 80% Lebar A4 (21cm) = 16.8cm
            p.add_run().add_picture(io.BytesIO(kop), width=Cm(16.8))
        except: pass
    
    # --- HEADER TAMBAHAN ---
    p_header = doc.add_paragraph()
    p_header.alignment = 1 # Center

    # Baris 1: LAPORAN KEGIATAN RHK X
    run1 = p_header.add_run(f"LAPORAN KEGIATAN {meta.get('rhk_id', 'RHK ...').upper()}\n")
    run1.bold = True
    run1.font.size = Pt(12)

    # Baris 2: Nama Kegiatan
    run2 = p_header.add_run(f"{meta.get('kegiatan_spesifik', 'Isi Pilihan Laporan Harian')}\n")
    run2.bold = True
    run2.font.size = Pt(11)

    # Baris 3: Bulan Tahun
    run3 = p_header.add_run(f"{meta['bulan'].upper()}")
    run3.bold = True
    run3.font.size = Pt(11)

    
    # Helper untuk paragraf isi
    def add_text_body(text, bold=False):
        p = doc.add_paragraph(str(text))
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if bold: p.runs[0].bold = True

    # --- A. PENDAHULUAN ---
    doc.add_paragraph("A. Pendahuluan", style='Heading 1')
    p_data = data.get('pendahuluan', {})
    
    doc.add_paragraph("    1. Umum", style='Normal')
    add_text_body(p_data.get('umum', '-'))
    
    doc.add_paragraph("    2. Maksud dan Tujuan", style='Normal')
    add_text_body(p_data.get('maksud_tujuan', '-'))

    doc.add_paragraph("    3. Ruang Lingkup", style='Normal')
    add_text_body(p_data.get('ruang_lingkup', '-'))

    doc.add_paragraph("    4. Dasar", style='Normal')
    dasar = p_data.get('dasar', [])
    if isinstance(dasar, list):
        for item in dasar: 
            # Filter Safety: Jika AI masih bandel mengeluarkan Surat Tugas
            if "Surat Tugas" not in item:
                doc.add_paragraph(str(item), style='List Bullet')
    else:
        add_text_body(str(dasar))
        
        # --- B. KEGIATAN YANG DILAKSANAKAN ---
    doc.add_paragraph("B. Kegiatan yang dilaksanakan", style='Heading 1')
    add_text_body(data.get('kegiatan', '-'))
        
    if kpm_data:
        doc.add_paragraph("    Data Peserta/KPM:", style='Normal')
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        for k, v in kpm_data.items(): 
            row = table.add_row().cells
            row[0].text = str(k)
            row[1].text = str(v)
        doc.add_paragraph("\n")

    # --- C. HASIL YANG DICAPAI ---
    doc.add_paragraph("C. Hasil yang dicapai", style='Heading 1')
    hasil = data.get('hasil', '-')
    if isinstance(hasil, list):
        for h in hasil: doc.add_paragraph(str(h), style='List Bullet')
    else:
        add_text_body(hasil)

    # --- D. SIMPULAN DAN SARAN ---
    doc.add_paragraph("D. Simpulan dan Saran", style='Heading 1')
    add_text_body(data.get('simpulan_saran', '-'))

    # --- E. PENUTUP ---
    doc.add_paragraph("E. Penutup", style='Heading 1')
    add_text_body(data.get('penutup', '-'))

    doc.add_paragraph("\n\n")
    
    # --- TANDA TANGAN ---
    table = doc.add_table(rows=1, cols=2); table.autofit = False
    table.columns[0].width = Inches(3); table.columns[1].width = Inches(3)
    c2 = table.cell(0, 1).paragraphs[0]; c2.alignment = 1
    
    c2.add_run(f"{meta['kab']}, {meta['tgl']}\n{meta.get('jabatan','Pendamping Sosial')}\n\n")
    if ttd: 
        try: c2.add_run().add_picture(io.BytesIO(ttd), height=Inches(0.8))
        except: pass
    c2.add_run(f"\n{meta['nama']}\nNIP. {meta['nip']}")
    
    # --- DOKUMENTASI (UPDATED) ---
    if imgs:
        doc.add_page_break()
        doc.add_paragraph("LAMPIRAN DOKUMENTASI", style='Heading 1').alignment = 1
        
        for i, img in enumerate(imgs):
            try: 
                # 1. Gambar (Posisi Tengah)
                p_img = doc.add_paragraph()
                p_img.alignment = 1 # 1 = Center
                # Lebar diatur 4.5 inci agar proporsional di A4
                p_img.add_run().add_picture(compress_image(img), width=Inches(4.5))
                
                # 2. Keterangan di BAWAH Foto (Italic, Font 9)
                p_cap = doc.add_paragraph()
                p_cap.alignment = 1 # Center
                run_cap = p_cap.add_run(f"Gambar {i+1}: Dokumentasi {meta.get('kegiatan_spesifik', 'Kegiatan')}")
                run_cap.italic = True
                run_cap.font.size = Pt(9) 
                
                # Spasi antar foto
                doc.add_paragraph("\n")
            except: pass
            
    bio = io.BytesIO(); doc.save(bio); return bio

def create_pdf_doc(data, meta, imgs, kop, ttd, extra_info=None, kpm_data=None):
    if data is None: return None
    pdf = FPDF('P', 'mm', 'A4')
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    # REVISI 1: Mengurangi margin kanan menjadi 15mm agar TTD bisa geser lebih kanan
    pdf.set_margins(25, 25, 15) 

    # --- KOP ---
    if kop:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
            tmp.write(kop); tmp.flush(); tmp_path = tmp.name
        try: 
            # --- FIX OVERLAPPING DENGAN MENGHITUNG TINGGI GAMBAR ---
            img_obj = Image.open(io.BytesIO(kop))
            w_px, h_px = img_obj.size
            
            # Lebar target di PDF: 168mm (80% dari 210mm)
            target_w = 210 * 0.8
            # Hitung tinggi proporsional di PDF (mm)
            target_h = (target_w / w_px) * h_px
            
            x_kop = (210 - target_w) / 2
            pdf.image(tmp_path, x=x_kop, y=10, w=target_w)
            
            # Set posisi Y di bawah gambar + margin 5mm agar tidak tertumpuk
            pdf.set_y(10 + target_h + 5) 
        except: pdf.ln(10)
        finally: 
            if os.path.exists(tmp_path): os.remove(tmp_path)
    else: pdf.ln(10)

    # --- HEADER TAMBAHAN ---
    pdf.set_font("Arial", "B", 12)
    # Baris 1: RHK ID
    pdf.cell(0, 6, f"LAPORAN KEGIATAN {clean_text_for_pdf(meta.get('rhk_id', 'RHK ...').upper())}", ln=True, align='C')
    # Baris 2: Kegiatan Spesifik
    pdf.set_font("Arial", "B", 11)
    pdf.multi_cell(0, 6, f"{clean_text_for_pdf(meta.get('kegiatan_spesifik', 'Isi Pilihan Laporan Harian'))}", align='C')
    # Baris 3: Bulan Tahun
    pdf.cell(0, 6, f"{clean_text_for_pdf(meta['bulan'].upper())}", ln=True, align='C')

    pdf.ln(8)
    
    def add_paragraph_pdf(text):
        pdf.set_font("Arial", "", 11)
        pdf.multi_cell(0, 6, clean_text_for_pdf(str(text)))
        pdf.ln(2)

    # --- A. PENDAHULUAN ---
    pdf.set_font("Arial", "B", 11); pdf.cell(0, 7, "A. Pendahuluan", ln=True); pdf.set_font("Arial", "", 11)
    p_data = data.get('pendahuluan', {})
    
    pdf.set_x(30); pdf.cell(0, 6, "1. Umum", ln=True)
    pdf.set_x(35); pdf.multi_cell(0, 6, clean_text_for_pdf(p_data.get('umum', '-')))
    
    pdf.set_x(30); pdf.cell(0, 6, "2. Maksud dan Tujuan", ln=True)
    pdf.set_x(35); pdf.multi_cell(0, 6, clean_text_for_pdf(p_data.get('maksud_tujuan', '-')))

    pdf.set_x(30); pdf.cell(0, 6, "3. Ruang Lingkup", ln=True)
    pdf.set_x(35); pdf.multi_cell(0, 6, clean_text_for_pdf(p_data.get('ruang_lingkup', '-')))

    pdf.set_x(30); pdf.cell(0, 6, "4. Dasar", ln=True)
    dasar = p_data.get('dasar', [])
    if isinstance(dasar, list):
        for item in dasar: 
            if "Surat Tugas" not in item:
                pdf.set_x(35); pdf.multi_cell(0, 6, f"- {clean_text_for_pdf(item)}")
    else:
        pdf.set_x(35); pdf.multi_cell(0, 6, clean_text_for_pdf(str(dasar)))
    pdf.ln(3)

    # --- B. KEGIATAN ---
    pdf.set_font("Arial", "B", 11); pdf.cell(0, 7, "B. Kegiatan yang dilaksanakan", ln=True); pdf.set_font("Arial", "", 11)
    add_paragraph_pdf(data.get('kegiatan', '-'))
    
    if kpm_data:
        pdf.ln(2)
        pdf.set_font("Arial", "B", 10); pdf.cell(0, 6, "Data Peserta:", ln=True); pdf.set_font("Arial", "", 10)
        col_w = 80
        for k, v in kpm_data.items(): 
            pdf.cell(col_w, 6, clean_text_for_pdf(str(k)), border=1)
            pdf.cell(col_w, 6, clean_text_for_pdf(str(v)), border=1, ln=True)
        pdf.ln(3)

    # --- C. HASIL ---
    pdf.set_font("Arial", "B", 11); pdf.cell(0, 7, "C. Hasil yang dicapai", ln=True); pdf.set_font("Arial", "", 11)
    hasil = data.get('hasil', '-')
    if isinstance(hasil, list):
        for h in hasil: 
            pdf.set_x(30); pdf.multi_cell(0, 6, f"- {clean_text_for_pdf(h)}")
    else:
        add_paragraph_pdf(hasil)
    pdf.ln(3)

    # --- D. SIMPULAN SARAN ---
    pdf.set_font("Arial", "B", 11); pdf.cell(0, 7, "D. Simpulan dan Saran", ln=True); pdf.set_font("Arial", "", 11)
    add_paragraph_pdf(data.get('simpulan_saran', '-'))
    pdf.ln(3)

    # --- E. PENUTUP ---
    pdf.set_font("Arial", "B", 11); pdf.cell(0, 7, "E. Penutup", ln=True); pdf.set_font("Arial", "", 11)
    add_paragraph_pdf(data.get('penutup', '-'))
    pdf.ln(10)

    # --- TTD ---
    if pdf.get_y() > 220: pdf.add_page()
    # REVISI 2: Geser X ke 140 (lebih kanan) dan perkecil W ke 55
    x_block = 140; w_block = 55
    pdf.set_x(x_block); pdf.multi_cell(w_block, 6, f"{clean_text_for_pdf(meta['kab'])}, {clean_text_for_pdf(meta['tgl'])}", align='C')
    pdf.set_x(x_block); pdf.multi_cell(w_block, 6, clean_text_for_pdf(meta.get('jabatan', 'Pendamping Sosial')), align='C')
    
    if ttd:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp_ttd: 
            tmp_ttd.write(ttd); tmp_ttd.flush(); ttd_path = tmp_ttd.name
        try: 
            y_img = pdf.get_y()
            pdf.image(ttd_path, x=x_block + 5, y=y_img, h=25)
            pdf.set_y(y_img + 27)
        except: pdf.ln(25)
        finally: 
            if os.path.exists(ttd_path): os.remove(ttd_path)
    else: pdf.ln(25)
    
    pdf.set_x(x_block); pdf.set_font("Arial", "BU", 11); pdf.cell(w_block, 6, clean_text_for_pdf(meta['nama']), ln=True, align='C')
    pdf.set_x(x_block); pdf.set_font("Arial", "", 11); pdf.cell(w_block, 6, f"NIP. {clean_text_for_pdf(meta['nip'])}", ln=True, align='C')

    # --- LAMPIRAN ---
    if imgs:
        pdf.add_page()
        pdf.set_font("Arial", "B", 12); pdf.cell(0, 10, "LAMPIRAN DOKUMENTASI", ln=True, align='C'); pdf.ln(5)
        for img_bytes in imgs:
            compressed = compress_image(img_bytes)
            with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as tmp_img: 
                tmp_img.write(compressed.getvalue()); tmp_img.flush(); img_path = tmp_img.name
            try: 
                x_center = (210 - 120) / 2
                pdf.image(img_path, x=x_center, w=120)
                pdf.ln(5)
            except: pass
            finally: 
                if os.path.exists(img_path): os.remove(img_path)
    return pdf.output(dest='S').encode('latin-1')

# ==========================================
# 5. UI & LOGIC UTAMA (MULTI AKUN)
# ==========================================

def check_password():
    if st.session_state.get("password_correct", False): return True
    
    st.markdown("<br><h1 style='text-align: center;'>🔐 APLIKASI RHK PRO 2.0</h1>", unsafe_allow_html=True)
    
    tab1, tab2 = st.tabs(["MASUK (LOGIN)", "DAFTAR BARU"])
    
    # --- TAB LOGIN ---
    with tab1:
        with st.form("login_form"):
            input_user = st.text_input("Username")
            input_pass = st.text_input("Password", type="password")
            if st.form_submit_button("MASUK", type="primary", use_container_width=True):
                # 1. Cek Default/Admin
                if input_user in DAFTAR_USER and DAFTAR_USER[input_user] == input_pass:
                    st.session_state["password_correct"] = True
                    st.session_state["username"] = input_user
                    st.rerun()
                # 2. Cek Database Users
                else:
                    conn = sqlite3.connect('rhk_pro_fixed.db'); c = conn.cursor()
                    c.execute("SELECT password FROM users WHERE username=?", (input_user,))
                    data = c.fetchone()
                    conn.close()
                    
                    if data and data[0] == input_pass:
                        st.session_state["password_correct"] = True
                        st.session_state["username"] = input_user
                        st.rerun()
                    else:
                        st.error("😕 Username atau Password Salah!")

    # --- TAB REGISTRASI ---
    with tab2:
        with st.form("reg_form"):
            st.write("Buat Akun Baru")
            new_user = st.text_input("Username Baru")
            new_pass = st.text_input("Password Baru", type="password")
            if st.form_submit_button("DAFTAR", type="secondary", use_container_width=True):
                if new_user and new_pass:
                    conn = sqlite3.connect('rhk_pro_fixed.db'); c = conn.cursor()
                    try:
                        c.execute("INSERT INTO users (username, password) VALUES (?, ?)", (new_user, new_pass))
                        conn.commit()
                        st.success("✅ Akun berhasil dibuat! Silakan login di tab sebelah.")
                    except sqlite3.IntegrityError:
                        st.error("❌ Username sudah terpakai.")
                    finally:
                        conn.close()
                else:
                    st.warning("Isi username dan password.")

    return False

if check_password():
    init_db()
    
    # Ambil user yang sedang login
    current_user = st.session_state.get('username')
    
    keys = ['page', 'selected_rhk', 'kop_bytes', 'ttd_bytes', 'rhk2_queue', 'rhk2_results', 'rhk3_queue', 'rhk3_results', 'rhk4_graduasi_results', 'rhk8_queue', 'rhk8_results', 'generated_file_data', 'tgl_val', 'bln_val', 'th_val'] 
    for k in keys:
        if k not in st.session_state: st.session_state[k] = None
    if st.session_state['rhk2_queue'] is None: st.session_state['rhk2_queue'] = []
    if st.session_state['rhk3_queue'] is None: st.session_state['rhk3_queue'] = []
    if st.session_state['rhk8_queue'] is None: st.session_state['rhk8_queue'] = []
    if st.session_state['page'] is None: st.session_state['page'] = 'home'
    if not st.session_state['bln_val']: st.session_state['bln_val'] = "JANUARI"
    if not st.session_state['th_val']: st.session_state['th_val'] = "2026"
    if not st.session_state['tgl_val']: st.session_state['tgl_val'] = "30 Januari 2026"

    def update_tanggal(): st.session_state.tgl_val = f"30 {st.session_state.bln_val.title()} {st.session_state.th_val}"
    
    # Load settings based on current user
    u_nama, u_nip, u_kpm, u_prov, u_kab, u_kec, u_kel, u_jabatan = get_user_settings(current_user)
    
    with st.sidebar:
        st.write(f"👤 User: **{current_user}**")
        if st.button("🔒 Logout", type="primary"): st.session_state["password_correct"] = False; st.query_params.clear(); st.rerun()
        with st.expander("👤 Profil", expanded=False):
            with st.form("profil_form"):
                nama = st.text_input("Nama", u_nama); nip = st.text_input("NIP", u_nip)
                jabatan = st.text_input("Jabatan", value=u_jabatan if u_jabatan else "Pendamping Sosial")
                kpm = st.number_input("Jml KPM", value=u_kpm); prov = st.text_input("Provinsi", u_prov)
                kab = st.text_input("Kabupaten", u_kab); kec = st.text_input("Kecamatan", u_kec); kel = st.text_input("Kelurahan", u_kel)
                if st.form_submit_button("Simpan Profil"): 
                    save_user_settings(current_user, nama, nip, kpm, prov, kab, kec, kel, jabatan)
                    st.success("Tersimpan!")
                    st.rerun()
        st.markdown("---")
        st.selectbox("Bulan", ["JANUARI", "FEBRUARI", "MARET", "APRIL", "MEI", "JUNI", "JULI", "AGUSTUS", "SEPTEMBER", "OKTOBER", "NOVEMBER", "DESEMBER"], key="bln_val", on_change=update_tanggal)
        st.selectbox("Tahun", ["2026", "2027"], key="th_val", on_change=update_tanggal)
        st.text_input("Tanggal Surat", key="tgl_val")
        st.markdown("---")
        if st.button("🗄️ ARSIP LAPORAN", use_container_width=True, type="secondary"): st.session_state['page'] = 'history'; st.rerun()
        st.markdown("---")
        st.file_uploader("Kop Surat (JPG/PNG)", type=['png','jpg'], key="kop_up")
        if st.session_state.get('kop_up'): st.session_state['kop_bytes'] = st.session_state['kop_up'].getvalue()
        st.file_uploader("Tanda Tangan (JPG/PNG)", type=['png','jpg'], key="ttd_up")
        if st.session_state.get('ttd_up'): st.session_state['ttd_bytes'] = st.session_state['ttd_up'].getvalue()
        
        # --- FOOTER CREATED BY ---
        st.markdown("---")
        st.markdown("""<div style='text-align: center; color: grey; font-size: 11px;'>Created by:<br><b>[VDStudio]</b><br>&copy; 2026 RHK PKH Pro</div>""", unsafe_allow_html=True)

    def show_dashboard():
        st.title("📂 Aplikasi RHK PKH Pro 2.0"); cols = st.columns(3)
        for i, rhk in enumerate(CONFIG_LAPORAN.keys()):
            with cols[i % 3]:
                st.markdown(f"""<div style="background-color:#f0f2f6; padding:15px; border-radius:10px; margin-bottom:10px; border:1px solid #d1d5db;"><b>{rhk.split('–')[0]}</b><br><small>{rhk.split('–')[-1]}</small></div>""", unsafe_allow_html=True)
                if st.button(f"Buka {rhk.split('–')[0]}", key=f"nav_{i}", use_container_width=True): st.session_state['selected_rhk'] = rhk; st.session_state['page'] = 'detail'; st.rerun()
        st.markdown("---"); 
        if st.button("🗄️ Buka Arsip Digital (History Laporan)", use_container_width=True): st.session_state['page'] = 'history'; st.rerun()

    def show_history_page():
        st.title("🗄️ Arsip Digital Laporan")
        col1, col2, col3 = st.columns([2, 2, 4])
        with col1: f_bulan = st.selectbox("Filter Bulan", ["SEMUA", "JANUARI", "FEBRUARI", "MARET", "APRIL", "MEI", "JUNI", "JULI", "AGUSTUS", "SEPTEMBER", "OKTOBER", "NOVEMBER", "DESEMBER"])
        with col2: f_tahun = st.selectbox("Filter Tahun", ["SEMUA", "2026", "2027"])
        with col3: 
            if st.button("⬅️ Kembali ke Dashboard", use_container_width=True): st.session_state['page'] = 'home'; st.rerun()
        
        # Load history based on current user
        data = get_history_by_filter(current_user, f_bulan, f_tahun)
        
        if not data: st.info("Belum ada riwayat laporan yang tersimpan."); return
        st.markdown(f"**Ditemukan: {len(data)} Dokumen**")
        for row in data:
            with st.expander(f"{row['jenis_rhk']} - {row['judul_kegiatan']} ({row['bulan']} {row['tahun']})"):
                c1, c2, c3, c4 = st.columns([2, 2, 2, 1])
                with c1: st.caption(f"📅 Dibuat: {row['tanggal']}")
                with c2: 
                    if st.button("📄 Word", key=f"wd_{row['id']}"): blob, title = get_file_from_history(row['id'], 'docx'); st.download_button(label="⬇️ Download Docx", data=blob, file_name=f"{title}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"dwd_real_{row['id']}")
                with c3:
                    if st.button("📕 PDF", key=f"pdf_{row['id']}"): blob, title = get_file_from_history(row['id'], 'pdf'); st.download_button(label="⬇️ Download PDF", data=blob, file_name=f"{title}.pdf", mime="application/pdf", key=f"dpdf_real_{row['id']}")
                with c4:
                    if st.button("🗑️", key=f"del_h_{row['id']}"): delete_history_item(row['id']); st.success("Terhapus"); time.sleep(0.5); st.rerun()

    def show_detail():
        rhk = st.session_state.get('selected_rhk')
        if rhk not in CONFIG_LAPORAN: st.warning("🔄 Refreshing session..."); st.session_state['page'] = 'home'; st.rerun(); return
        c1, c2 = st.columns([1, 6]); 
        if c1.button("⬅️ Kembali", use_container_width=True): st.session_state['page'] = 'home'; st.rerun()
        c2.markdown(f"### 📝 {rhk}")
        
        # Update Meta Data: Tambahkan rhk_id untuk header
        meta = {
            'bulan': f"{st.session_state.bln_val} {st.session_state.th_val}", 
            'nama': u_nama, 
            'nip': u_nip, 
            'jabatan': u_jabatan if u_jabatan else "Pendamping Sosial", 
            'kab': u_kab, 
            'kec': u_kec, 
            'kel': u_kel, 
            'tgl': st.session_state.tgl_val, 
            'judul': rhk.split('–')[-1].upper(),
            'rhk_id': rhk.split('–')[0] # Contoh: RHK 1
        }
        
        lokasi = f"{u_kel}, {u_kec}, {u_kab}"
        if "RHK 2" in rhk or "RHK 3" in rhk or "RHK 8" in rhk:
            q_key = 'rhk2_queue' if "RHK 2" in rhk else ('rhk3_queue' if "RHK 3" in rhk else 'rhk8_queue'); r_key = q_key.replace('queue', 'results')
            st.info("💡 **Mode Antrian:** Masukkan kegiatan satu per satu, lalu klik 'Generate Semua'.")
            DATA_P2K2 = {"Modul Ekonomi": ["1. Mengelola Keuangan Keluarga", "2. Cermat Meminjam dan Menabung", "3. Memulai Usaha"], "Modul Kesehatan dan Gizi": ["1. Pentingnya Gizi dan Layanan Ibu Hamil", "2. Pentingnya Gizi Untuk Ibu Menyusui dan Balita", "3. Kesakitan pada Anak dan Kesehatan Lingkungan"], "Modul Kesejahteraan": ["1. Pelayanan bagi Penyandang Disabilitas Berat", "2. Pentingnya Kesejahteraan Lanjut Usia"], "Modul Pengasuhan dan Pendidikan": ["1. Menjadi Orang Tua yang Lebih Baik", "2. Memahami Perilaku Anak", "3. Memahami Cara Anak Usia Dini Belajar", "4. Membantu Anak Sukses di Sekolah"], "Modul Perlindungan Anak": ["1. Pencegahan Kekerasan Terhadap Anak", "2. Penelantaran dan Eksploitasi Anak"], "Modul Stunting": ["1. Pencegahan dan Penanganan Stunting", "2. Dukungan Pemenuhan Kesejahteraan Bayi Baru Lahir dan Ibu Menyusui"]}
            if "RHK 2" in rhk:
                kegiatan = st.selectbox("Sub-Kegiatan", CONFIG_LAPORAN[rhk])
                col_m1, col_m2 = st.columns(2)
                with col_m1: pilih_modul = st.selectbox("Pilih Modul P2K2", list(DATA_P2K2.keys()), key="p2k2_modul_sel")
                with col_m2: opsi_sesi = DATA_P2K2.get(pilih_modul, ["-"]); pilih_sesi = st.selectbox("Pilih Sesi", opsi_sesi, key="p2k2_sesi_sel")
                with st.form("queue_form_rhk2", clear_on_submit=True):
                    ket_q = st.text_input("Keterangan Tambahan", placeholder="Nama Kelompok / Lokasi / Jumlah Peserta...")
                    fotos = st.file_uploader("Foto Kegiatan", accept_multiple_files=True, type=['jpg','png'])
                    if st.form_submit_button("➕ Tambah ke Antrian"):
                        if not fotos: st.error("❌ Foto wajib diupload!")
                        else: final_judul = f"{kegiatan} - {pilih_modul}"; final_ket = f"Materi: {pilih_sesi}. Keterangan: {ket_q}"; st.session_state[q_key].append({"kegiatan": final_judul, "ket": final_ket, "fotos": [io.BytesIO(f.getvalue()) for f in fotos]}); st.success(f"Berhasil: {pilih_sesi} ditambahkan!"); st.rerun()
            else:
                with st.form("queue_form", clear_on_submit=True):
                    try: kegiatan = st.selectbox("Sub-Kegiatan", CONFIG_LAPORAN[rhk])
                    except: kegiatan = "Kegiatan Umum"
                    ket_q = st.text_input("Keterangan", placeholder="Detail lokasi/peserta...")
                    fotos = st.file_uploader("Foto", accept_multiple_files=True, type=['jpg','png'])
                    if st.form_submit_button("➕ Tambah"):
                        if not fotos: st.error("❌ Foto wajib!")
                        else: st.session_state[q_key].append({"kegiatan": kegiatan, "ket": ket_q, "fotos": [io.BytesIO(f.getvalue()) for f in fotos]}); st.success("Masuk antrian!"); st.rerun()
            queue = st.session_state[q_key]
            if queue:
                st.write(f"**Antrian ({len(queue)} Item):**")
                for i, q in enumerate(queue):
                    col_txt, col_del = st.columns([0.85, 0.15])
                    with col_txt: st.text(f"{i+1}. {q['kegiatan']} - {q['ket']}")
                    with col_del: 
                        if st.button("🗑️ Hapus", key=f"del_{q_key}_{i}"): st.session_state[q_key].pop(i); st.rerun()
                if st.button("🚀 GENERATE SEMUA & SIMPAN ARSIP", type="primary"):
                    results = []; bar = st.progress(0)
                    with st.spinner("Sedang menghubungi AI dan menyusun dokumen..."):
                        for i, item in enumerate(queue):
                            jd = generate_isi_laporan(rhk, item['kegiatan'], u_kpm, "Peserta", meta['bulan'], lokasi, item['ket'])
                            # Update meta untuk kegiatan spesifik saat generate doc
                            meta['kegiatan_spesifik'] = item['kegiatan']
                            if jd: 
                                w = create_word_doc(jd, meta, item['fotos'], st.session_state['kop_bytes'], st.session_state['ttd_bytes'], item['ket'])
                                p = create_pdf_doc(jd, meta, item['fotos'], st.session_state['kop_bytes'], st.session_state['ttd_bytes'], item['ket'])
                                if w: 
                                    results.append({"judul": item['kegiatan'], "file": w, "file_pdf": p})
                                    save_to_history(current_user, meta['bulan'], st.session_state.th_val, rhk.split('–')[0], item['kegiatan'], w, p)
                            bar.progress((i + 1) / len(queue))
                    st.session_state[r_key] = results; st.success("Selesai! Laporan telah tersimpan di Arsip."); st.rerun()
            if st.session_state.get(r_key):
                st.write("### 📥 Download Hasil")
                for i, r in enumerate(st.session_state[r_key]):
                    c1, c2 = st.columns(2)
                    with c1: st.download_button(f"📄 Word: {r['judul']}", r['file'], f"{r['judul']}.docx", key=f"dw{i}", use_container_width=True)
                    with c2: 
                        if r.get('file_pdf'): st.download_button(f"📕 PDF: {r['judul']}", r['file_pdf'], f"{r['judul']}.pdf", key=f"dp{i}", use_container_width=True)
        elif "RHK 4" in rhk:
            st.info("ℹ️ **Mode Graduasi:** Upload Excel KPM."); df_tmpl = pd.DataFrame({"Nama": ["Budi"], "NIK": ["123"], "Alamat": ["Desa A"], "Kategori": ["PKH"], "Status": ["Graduasi"], "Alasan": ["Mampu"]}); buf = io.BytesIO(); df_tmpl.to_excel(buf, index=False); buf.seek(0); st.download_button("📥 Template Excel", buf, "Template.xlsx")
            upl = st.file_uploader("Upload Excel", type=['xlsx'])
            if upl:
                try:
                    df = pd.read_excel(upl); sel_kpm = st.multiselect("Pilih KPM", df['Nama'].tolist()) if 'Nama' in df.columns else []
                    if sel_kpm:
                        photos = st.file_uploader("Foto", accept_multiple_files=True)
                        if st.button("🚀 Generate & Simpan") and photos:
                            res = []; p_data = [io.BytesIO(f.getvalue()) for f in photos]; bar = st.progress(0)
                            with st.spinner("Memproses data graduasi..."):
                                for i, nm in enumerate(sel_kpm):
                                    row = df[df['Nama'] == nm].iloc[0].to_dict(); jd = generate_isi_laporan(rhk, f"Graduasi {nm}", 1, nm, meta['bulan'], lokasi, f"Graduasi {nm}")
                                    meta['kegiatan_spesifik'] = f"Graduasi KPM: {nm}"
                                    if jd:
                                        w = create_word_doc(jd, meta, p_data, st.session_state['kop_bytes'], st.session_state['ttd_bytes'], f"Graduasi {nm}", row)
                                        p = create_pdf_doc(jd, meta, p_data, st.session_state['kop_bytes'], st.session_state['ttd_bytes'], f"Graduasi {nm}", row)
                                        if w: 
                                            res.append({"judul": nm, "file": w, "file_pdf": p})
                                            save_to_history(current_user, meta['bulan'], st.session_state.th_val, rhk.split('–')[0], nm, w, p)
                                    bar.progress((i+1)/len(sel_kpm))
                            st.session_state['rhk4_graduasi_results'] = res; st.success("Selesai! Laporan tersimpan di Arsip."); st.rerun()
                except: st.error("Format Excel Salah")
            if st.session_state.get('rhk4_graduasi_results'):
                for i, r in enumerate(st.session_state['rhk4_graduasi_results']):
                    c1, c2 = st.columns(2)
                    with c1: st.download_button(f"📥 Word: {r['judul']}", r['file'], f"Graduasi_{r['judul']}.docx", key=f"dgw{i}", use_container_width=True)
                    with c2: 
                        if r.get('file_pdf'): st.download_button(f"📥 PDF: {r['judul']}", r['file_pdf'], f"Graduasi_{r['judul']}.pdf", key=f"dgp{i}", use_container_width=True)
        elif "RHK 5" in rhk:
            st.info("ℹ️ **Mode Verivali:** Pilih aplikasi yang digunakan untuk verifikasi.")
            LIST_APPS = ["SIKS-NG (Sistem Informasi Kesejahteraan Sosial - Next Generation)", "SIKS-Mobile (Android)", "Aplikasi Cek Bansos", "DTKS Offline / Excel", "Verifikasi Manual / Berkas Fisik"]
            with st.form("rhk5_form"):
                kegiatan = st.selectbox("Sub-Kegiatan", CONFIG_LAPORAN[rhk]); aplikasi = st.selectbox("Aplikasi / Media", LIST_APPS); ket = st.text_area("Keterangan Tambahan", placeholder="Contoh: Perbaikan data anomali rekening..."); fotos = st.file_uploader("Foto Bukti (Screenshot Aplikasi/Lapangan)", accept_multiple_files=True)
                if st.form_submit_button("🚀 BUAT LAPORAN VERIVALI", type="primary"):
                    if not fotos: st.error("Foto wajib!")
                    else:
                        full_context = f"Menggunakan {aplikasi}. {ket}"
                        with st.spinner("Mengolah data Verivali..."):
                            jd = generate_isi_laporan(rhk, kegiatan, u_kpm, "KPM", meta['bulan'], lokasi, full_context)
                            meta['kegiatan_spesifik'] = kegiatan
                            if jd:
                                imgs_data = [io.BytesIO(f.getvalue()) for f in fotos]
                                w = create_word_doc(jd, meta, imgs_data, st.session_state['kop_bytes'], st.session_state['ttd_bytes'], full_context)
                                p = create_pdf_doc(jd, meta, imgs_data, st.session_state['kop_bytes'], st.session_state['ttd_bytes'], full_context)
                                st.session_state['generated_file_data'] = {"name": f"Laporan {kegiatan}", "file": w, "file_pdf": p}
                                save_to_history(current_user, meta['bulan'], st.session_state.th_val, rhk.split('–')[0], kegiatan, w, p)
                                st.success("Selesai! Laporan tersimpan di Arsip."); st.rerun()
                            else: st.error("Gagal koneksi AI, coba lagi.")
            if st.session_state.get('generated_file_data'):
                f = st.session_state['generated_file_data']; st.success("Selesai!")
                c1, c2 = st.columns(2)
                with c1: st.download_button("📥 Download Word", f['file'], f"{f['name']}.docx", type="primary", use_container_width=True)
                with c2: 
                    if f.get('file_pdf'): st.download_button("📕 Download PDF", f['file_pdf'], f"{f['name']}.pdf", type="secondary", use_container_width=True)
        else:
            with st.form("std_form"):
                try: jk = st.selectbox("Kegiatan", CONFIG_LAPORAN[rhk])
                except: jk = "Kegiatan Umum"
                ka = st.text_area("Keterangan"); ft = st.file_uploader("Foto", accept_multiple_files=True)
                if st.form_submit_button("🚀 BUAT LAPORAN", type="primary"):
                    if not ft: st.error("Foto wajib!")
                    else:
                        with st.spinner("Sedang menghubungi AI dan menyusun dokumen..."):
                            jd = generate_isi_laporan(rhk, jk, u_kpm, "Peserta", meta['bulan'], lokasi, ka)
                            meta['kegiatan_spesifik'] = jk
                            if jd:
                                imgs_data = [io.BytesIO(f.getvalue()) for f in ft]
                                w = create_word_doc(jd, meta, imgs_data, st.session_state['kop_bytes'], st.session_state['ttd_bytes'], ka)
                                p = create_pdf_doc(jd, meta, imgs_data, st.session_state['kop_bytes'], st.session_state['ttd_bytes'], ka)
                                if w:
                                    st.session_state['generated_file_data'] = {"name": f"Laporan {jk}", "file": w, "file_pdf": p}
                                    save_to_history(current_user, meta['bulan'], st.session_state.th_val, rhk.split('–')[0], jk, w, p)
                                    st.success("Selesai! Laporan tersimpan di Arsip."); st.rerun()
                            else: st.error("Gagal koneksi AI, coba lagi.")
            if st.session_state.get('generated_file_data'):
                f = st.session_state['generated_file_data']; st.success("Selesai!")
                c1, c2 = st.columns(2)
                with c1: st.download_button("📥 Download Word", f['file'], f"{f['name']}.docx", type="primary", use_container_width=True)
                with c2: 
                    if f.get('file_pdf'): st.download_button("📕 Download PDF", f['file_pdf'], f"{f['name']}.pdf", type="secondary", use_container_width=True)

    if st.session_state['page'] == 'home': show_dashboard()
    elif st.session_state['page'] == 'history': show_history_page()
    else: show_detail()





